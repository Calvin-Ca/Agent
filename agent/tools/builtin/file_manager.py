"""Built-in document file-management tools."""

from __future__ import annotations

import re
from pathlib import Path

from loguru import logger

from agent.infra.config import get_settings
from agent.tools.base import BaseTool, ToolOutput
from app.db import minio as minio_db

_VIDEO_EXTS = {".mp4", ".avi", ".mov", ".mkv", ".wmv", ".flv", ".webm"}


class FileManagerTool(BaseTool):
    name = "file_manager"
    description = "Read or write text files under the workspace"
    input_schema = {
        "type": "object",
        "properties": {
            "path": {"type": "string"},
            "mode": {"type": "string"},
            "content": {"type": "string"},
        },
        "required": ["path", "mode"],
    }

    def execute(self, **kwargs) -> ToolOutput:
        path = Path(kwargs["path"])
        mode = kwargs.get("mode", "read")
        if mode == "read":
            return ToolOutput(success=True, data=path.read_text(encoding="utf-8"))
        if mode == "write":
            path.write_text(kwargs.get("content", ""), encoding="utf-8")
            return ToolOutput(success=True, data={"path": str(path), "written": True})
        return ToolOutput(success=False, error=f"Unsupported mode: {mode}")


def _strip_bold(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def export_to_docx(title: str, content: str, filename: str = "") -> str:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()
    title_paragraph = document.add_heading(title, level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for raw_line in content.strip().splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("---"):
            document.add_paragraph("─" * 40)
        elif line.startswith("- ") or line.startswith("* "):
            document.add_paragraph(_strip_bold(line[2:]), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            document.add_paragraph(_strip_bold(re.sub(r"^\d+\.\s", "", line)), style="List Number")
        else:
            document.add_paragraph(_strip_bold(line))

    settings = get_settings()
    if not filename:
        safe_title = re.sub(r"[^\w\u4e00-\u9fff\-]", "_", title)[:50]
        filename = f"{safe_title}.docx"

    output_path = settings.export_dir / filename
    document.save(str(output_path))
    logger.info("Exported docx {}", output_path)
    return str(output_path)


def export_to_markdown(title: str, content: str, filename: str = "") -> str:
    settings = get_settings()
    if not filename:
        safe_title = re.sub(r"[^\w\u4e00-\u9fff\-]", "_", title)[:50]
        filename = f"{safe_title}.md"

    output_path = settings.export_dir / filename
    output_path.write_text(f"# {title}\n\n{content}", encoding="utf-8")
    logger.info("Exported markdown {}", output_path)
    return str(output_path)


def _is_video(name: str) -> bool:
    return Path(name).suffix.lower() in _VIDEO_EXTS


class ExportDocxTool(BaseTool):
    name = "export.docx"
    description = "Export markdown report content to a .docx document"

    def execute(self, *, title: str, content: str, filename: str = "", **kwargs) -> ToolOutput:
        try:
            return ToolOutput(success=True, data=export_to_docx(title, content, filename))
        except Exception as exc:
            return ToolOutput(success=False, error=str(exc))


class ExportMarkdownTool(BaseTool):
    name = "export.markdown"
    description = "Export report content to a markdown file"

    def execute(self, *, title: str, content: str, filename: str = "", **kwargs) -> ToolOutput:
        try:
            return ToolOutput(success=True, data=export_to_markdown(title, content, filename))
        except Exception as exc:
            return ToolOutput(success=False, error=str(exc))


class GetLatestVideoTool(BaseTool):
    name = "minio.get_latest_video"
    description = "Scan MinIO and return the latest video metadata with a presigned URL"

    def execute(self, **kwargs) -> ToolOutput:
        try:
            objects = minio_db.list_objects(prefix="", recursive=True)
            videos = [obj for obj in objects if _is_video(obj.object_name)]
            if not videos:
                return ToolOutput(
                    success=True,
                    data=None,
                    metadata={"message": "MinIO bucket 中未找到任何视频文件"},
                )

            videos.sort(key=lambda item: Path(item.object_name).name, reverse=True)
            latest = videos[0]
            url = minio_db.presigned_get_url(latest.object_name, expires_seconds=3600)
            return ToolOutput(
                success=True,
                data={
                    "filename": Path(latest.object_name).name,
                    "object_key": latest.object_name,
                    "size": latest.size,
                    "last_modified": latest.last_modified.isoformat() if latest.last_modified else "",
                    "presigned_url": url,
                    "total_videos": len(videos),
                },
            )
        except Exception as exc:
            return ToolOutput(success=False, error=str(exc))
