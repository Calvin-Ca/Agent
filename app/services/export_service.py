"""Export service — convert reports to Word (.docx) or PDF."""

from __future__ import annotations

import re
from pathlib import Path

from loguru import logger

from app.config import get_settings


def export_to_docx(title: str, content: str, filename: str = "") -> str:
    """Export Markdown report content to a Word document.

    Args:
        title: Report title
        content: Report content in Markdown
        filename: Output filename (auto-generated if empty)

    Returns:
        File path of the generated .docx
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse Markdown content into docx elements
    lines = content.strip().splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        # Horizontal rule
        elif line.startswith("---"):
            doc.add_paragraph("─" * 40)
        # Bullet list
        elif line.startswith("- ") or line.startswith("* "):
            text = line[2:]
            text = _strip_bold(text)
            doc.add_paragraph(text, style="List Bullet")
        # Numbered list
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            text = _strip_bold(text)
            doc.add_paragraph(text, style="List Number")
        # Regular paragraph
        else:
            text = _strip_bold(line)
            doc.add_paragraph(text)

        i += 1

    # Save
    settings = get_settings()
    export_dir = settings.export_dir
    if not filename:
        safe_title = re.sub(r'[^\w\u4e00-\u9fff\-]', '_', title)[:50]
        filename = f"{safe_title}.docx"

    filepath = export_dir / filename
    doc.save(str(filepath))
    logger.info("Exported docx: {} ({} bytes)", filepath, filepath.stat().st_size)
    return str(filepath)


def export_to_markdown(title: str, content: str, filename: str = "") -> str:
    """Export report as Markdown file.

    Args:
        title: Report title
        content: Report content in Markdown
        filename: Output filename

    Returns:
        File path of the generated .md
    """
    settings = get_settings()
    export_dir = settings.export_dir

    if not filename:
        safe_title = re.sub(r'[^\w\u4e00-\u9fff\-]', '_', title)[:50]
        filename = f"{safe_title}.md"

    filepath = export_dir / filename
    full_content = f"# {title}\n\n{content}"
    filepath.write_text(full_content, encoding="utf-8")

    logger.info("Exported markdown: {}", filepath)
    return str(filepath)


def _strip_bold(text: str) -> str:
    """Remove Markdown bold markers **text** → text."""
    return re.sub(r'\*\*(.+?)\*\*', r'\1', text)